"""
WC-DC-PM-v1.0 Post-Processing Script
Applies mandatory formatting rules from Style Guide Section 12.
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

docx_path = 'WC-DC-PM-v1.0.docx'
print(f"Post-processing: {docx_path}")

doc = Document(docx_path)

# Count fixes applied
fixes = {
    'table_headers': 0,
    'widow_orphan': 0,
    'table_para_spacing': 0,
    'code_spacing': 0,
}

# 1. Repeat header row on ALL tables with 2+ rows
for table in doc.tables:
    if len(table.rows) >= 2:
        trPr = table.rows[0]._tr.get_or_add_trPr()
        if not trPr.findall(qn('w:tblHeader')):
            etree.SubElement(trPr, qn('w:tblHeader'))
            fixes['table_headers'] += 1

# 2. Widow/orphan control on all body paragraphs
for p in doc.paragraphs:
    if p.text.strip():
        pPr = p._element.get_or_add_pPr()
        for existing in pPr.findall(qn('w:widowControl')):
            pPr.remove(existing)
        wc = etree.SubElement(pPr, qn('w:widowControl'))
        wc.set(qn('w:val'), '1')
        fixes['widow_orphan'] += 1

# 3. Table -> Paragraph spacing (minimum 8pt gap after tables)
body = doc.element.body
children = list(body)
for i, child in enumerate(children):
    if child.tag.endswith('}tbl') or child.tag == 'tbl':
        for j in range(i + 1, min(i + 3, len(children))):
            next_tag = children[j].tag.split('}')[-1] if '}' in children[j].tag else children[j].tag
            if next_tag == 'p':
                p = Paragraph(children[j], doc)
                if p.text and p.text.strip():
                    sb = p.paragraph_format.space_before
                    if not sb or sb < Pt(8):
                        p.paragraph_format.space_before = Pt(8)
                        fixes['table_para_spacing'] += 1
                break

# 4. Code block spacing (Consolas font detection)
prev_code = False
for p in doc.paragraphs:
    if not p.runs:
        prev_code = False
        continue
    is_code = False
    for run in p.runs:
        if run.font.name and 'Consolas' in str(run.font.name):
            is_code = True
            break
    is_body = not is_code and p.text.strip()
    if is_code and not prev_code:
        sb = p.paragraph_format.space_before
        if not sb or sb < Pt(6):
            p.paragraph_format.space_before = Pt(6)
            fixes['code_spacing'] += 1
    if is_body and prev_code:
        sb = p.paragraph_format.space_before
        if not sb or sb < Pt(8):
            p.paragraph_format.space_before = Pt(8)
            fixes['code_spacing'] += 1
    prev_code = is_code

# Save
doc.save(docx_path)

print(f"\nPost-processing complete:")
print(f"  Table repeat headers added: {fixes['table_headers']}")
print(f"  Widow/orphan control set:   {fixes['widow_orphan']}")
print(f"  Table->Para spacing fixed:  {fixes['table_para_spacing']}")
print(f"  Code block spacing fixed:   {fixes['code_spacing']}")
print(f"\nSaved: {docx_path}")
