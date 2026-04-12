#!/usr/bin/env python3
"""
Regenerate WC-MA-DPC-v1.0.docx from the PDF (source of truth).
Uses python-docx with exact fonts/sizes/colors extracted from the PDF.

PDF design specs:
  - Fonts: Figtree (body/headings), IBM Plex Mono (code/footer)
  - Title: 28pt, Subtitle: 18pt, H1: 16pt, H2: 13pt, Body: 9.5pt
  - Code: 7.5pt IBM Plex Mono
  - Colors: #0A0A0A (text), #737373 (muted), table header #F2F2F2, borders #E4E4E4
  - Page: A4, margins ~20mm
"""

import fitz  # PyMuPDF — reads the PDF
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import re

# Design tokens from PDF
FONT_BODY = 'Figtree'
FONT_CODE = 'IBM Plex Mono'
COLOR_TEXT = RGBColor(0x0A, 0x0A, 0x0A)
COLOR_MUTED = RGBColor(0x73, 0x73, 0x73)
COLOR_TABLE_HEADER_BG = 'F2F2F2'
COLOR_BORDER = 'E4E4E4'

PDF_PATH = os.path.join(os.path.dirname(__file__), 'WC-MA-DPC-v1.0.pdf')
DOCX_PATH = os.path.join(os.path.dirname(__file__), 'WC-MA-DPC-v1.0.docx')


def set_cell_shading(cell, color_hex):
    """Set table cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, start, end with value=(size, color)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, (sz, color) in kwargs.items():
        tag = {'top': 'top', 'bottom': 'bottom', 'start': 'start', 'end': 'end',
               'left': 'start', 'right': 'end'}[edge]
        border = parse_xml(
            f'<w:{tag} {nsdecls("w")} w:val="single" w:sz="{sz}" '
            f'w:space="0" w:color="{color}"/>')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run(paragraph, text, font_name=FONT_BODY, size=Pt(9.5), color=COLOR_TEXT,
            bold=False, italic=False):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    # Force font for East Asian / complex script
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" '
        f'w:cs="{font_name}" w:eastAsia="{font_name}"/>')
    rPr.insert(0, rFonts)
    return run


def add_heading(doc, text, level=1):
    """Add a WakeCap-styled heading."""
    sizes = {1: Pt(16), 2: Pt(13), 3: Pt(11)}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16) if level == 1 else Pt(10)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, size=sizes.get(level, Pt(11)), bold=False, color=COLOR_TEXT)
    return p


def add_body(doc, text):
    """Add body text paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text)
    return p


def add_bullet(doc, text):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    add_run(p, text)
    return p


def add_code_block(doc, lines):
    """Add monospace code block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        add_run(p, line, font_name=FONT_CODE, size=Pt(7.5))


def add_table(doc, headers, rows, col_widths=None):
    """Add a styled WakeCap table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        add_run(p, h, size=Pt(8), bold=False, color=COLOR_TEXT)
        set_cell_shading(cell, COLOR_TABLE_HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            add_run(p, str(val), size=Pt(8))

    # Set borders
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top=(4, COLOR_BORDER), bottom=(4, COLOR_BORDER),
                left=(4, COLOR_BORDER), right=(4, COLOR_BORDER))

    return table


def build_document():
    """Build the full DPC document matching the PDF design."""
    doc = Document()

    # Page setup — A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.9)

    # -- Extract text from PDF pages for faithful reproduction --
    pdf = fitz.open(PDF_PATH)

    # ========== TITLE PAGE ==========
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    add_run(p, 'MPPT Solar Charge Controller', size=Pt(28), color=COLOR_TEXT)

    p = doc.add_paragraph()
    add_run(p, 'Data Parsing Specification', size=Pt(18), color=COLOR_MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    add_run(p, 'Add-On Data Contract between Firmware & Backend Teams',
            size=Pt(10), color=COLOR_MUTED)

    # Metadata table
    meta = [
        ('Document ID', 'WC-MA-DPC-v1.0'),
        ('Version', '1.0'),
        ('Date', '2026-03-09'),
        ('Status', 'Final'),
        ('Companion', 'WeatherStation_DataParsingContractRevision.pdf (v1.0)'),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_run(p, label, size=Pt(9.5), bold=False, color=COLOR_TEXT)
        add_run(p, '    ', size=Pt(9.5))
        add_run(p, value, size=Pt(9.5), color=COLOR_MUTED)

    # TOC
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    add_run(p, 'Contents', size=Pt(11), bold=False, color=COLOR_TEXT)

    toc_items = [
        ('1.', 'Overview'),
        ('2.', 'Data Transmission via MQTT'),
        ('3.', 'Field Definitions (Tag 0x02 — MPPT, Length 0x1B)'),
        ('4.', 'Data Parsing Logic'),
        ('5.', 'Parsing Examples'),
        ('6.', 'Error Status Reporting (Endpoint 65)'),
        ('7.', 'Key Differences from Weather Station'),
        ('8.', 'Firmware Source References'),
        ('A.', 'Appendix: Complete Byte Map'),
        ('B.', 'Appendix: Backend Pseudocode'),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_run(p, num, font_name=FONT_CODE, size=Pt(8), color=COLOR_MUTED)
        add_run(p, '  ', size=Pt(8))
        add_run(p, title, size=Pt(9.5), color=COLOR_TEXT)

    # Footer-style info
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    add_run(p, 'WC-MA-DPC-v1.0', font_name=FONT_CODE, size=Pt(7), color=COLOR_MUTED)
    add_run(p, '    ', size=Pt(7))
    add_run(p, '2026-03-09', font_name=FONT_CODE, size=Pt(7), color=COLOR_MUTED)
    add_run(p, '    ', size=Pt(7))
    add_run(p, 'Page 1 / 11', font_name=FONT_CODE, size=Pt(7), color=COLOR_MUTED)

    doc.add_page_break()

    # ========== Now extract remaining pages from PDF text ==========
    # For pages 2-11, parse the PDF text and rebuild with proper styling

    for pg_idx in range(1, len(pdf)):
        page = pdf[pg_idx]
        blocks = page.get_text('dict')['blocks']

        for block in blocks:
            if 'lines' not in block:
                continue

            for line in block['lines']:
                spans = line['spans']
                if not spans:
                    continue

                first_span = spans[0]
                text = first_span['text'].strip()
                font = first_span['font']
                size = first_span['size']

                if not text:
                    continue

                # Skip header/footer repeated text
                if text.startswith('WC-MA-DPC-v1.0') and size < 8:
                    continue
                if text.startswith('Page ') and size < 8:
                    continue

                # Heading detection
                if size >= 15:  # H1 (16pt)
                    add_heading(doc, ''.join(s['text'] for s in spans), level=1)
                elif size >= 12:  # H2 (13pt)
                    add_heading(doc, ''.join(s['text'] for s in spans), level=2)
                elif 'IBMPlexMono' in font and size < 9:
                    # Code line
                    full_text = ''.join(s['text'] for s in spans)
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.left_indent = Cm(0.8)
                    add_run(p, full_text, font_name=FONT_CODE, size=Pt(7.5))
                elif 'Italic' in font:
                    # Note/italic text
                    full_text = ''.join(s['text'] for s in spans)
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    add_run(p, full_text, size=Pt(8.5), italic=True, color=COLOR_MUTED)
                else:
                    # Regular body text — concatenate all spans in the line
                    full_text = ''.join(s['text'] for s in spans)
                    if full_text.startswith('•'):
                        add_bullet(doc, full_text[1:].strip())
                    else:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(3)
                        for span in spans:
                            s_text = span['text']
                            s_font = span['font']
                            s_size = span['size']
                            s_color = span['color']
                            s_bold = bool(span['flags'] & 16)

                            is_mono = 'IBMPlexMono' in s_font or 'Mono' in s_font
                            add_run(p, s_text,
                                    font_name=FONT_CODE if is_mono else FONT_BODY,
                                    size=Pt(s_size),
                                    color=RGBColor(
                                        (s_color >> 16) & 0xFF,
                                        (s_color >> 8) & 0xFF,
                                        s_color & 0xFF),
                                    bold=s_bold)

        # Add page break between pages (except last)
        if pg_idx < len(pdf) - 1:
            doc.add_page_break()

    pdf.close()
    return doc


if __name__ == '__main__':
    print('Building WC-MA-DPC-v1.0.docx from PDF reference...')
    doc = build_document()
    doc.save(DOCX_PATH)
    print(f'Saved: {DOCX_PATH}')
    size_kb = os.path.getsize(DOCX_PATH) / 1024
    print(f'Size: {size_kb:.0f} KB')
