"""Apply Style Guide Section 12 post-processing to WC-GW-QR-LED-v1.0.docx."""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree

DOCX = Path(__file__).resolve().parents[2] / "output" / "gateway" / "WC-GW-QR-LED-v1.0.docx"

doc = Document(str(DOCX))

# 1. Repeat header on every table with 2+ rows
for table in doc.tables:
    if len(table.rows) >= 2:
        trPr = table.rows[0]._tr.get_or_add_trPr()
        if not trPr.findall(qn("w:tblHeader")):
            etree.SubElement(trPr, qn("w:tblHeader"))

# 2. Widow/orphan control on body paragraphs
for p in doc.paragraphs:
    if p.text.strip():
        pPr = p._element.get_or_add_pPr()
        for existing in pPr.findall(qn("w:widowControl")):
            pPr.remove(existing)
        wc = etree.SubElement(pPr, qn("w:widowControl"))
        wc.set(qn("w:val"), "1")

# 3. Table -> Paragraph spacing (8pt before paragraph immediately following a table)
body = doc.element.body
children = list(body)
for i, child in enumerate(children):
    if child.tag.split("}")[-1] == "tbl":
        for j in range(i + 1, min(i + 3, len(children))):
            if children[j].tag.split("}")[-1] == "p":
                p = Paragraph(children[j], doc)
                if p.text.strip():
                    sb = p.paragraph_format.space_before
                    if not sb or sb < Pt(8):
                        p.paragraph_format.space_before = Pt(8)
                break

# 4. Code block spacing
prev_code = False
for p in doc.paragraphs:
    if not p.runs:
        prev_code = False
        continue
    is_code = bool(p.runs[0].font.name and "Consolas" in str(p.runs[0].font.name))
    is_body = (not is_code) and bool(p.text.strip())
    if is_code and not prev_code:
        sb = p.paragraph_format.space_before
        if not sb or sb < Pt(6):
            p.paragraph_format.space_before = Pt(6)
    if is_body and prev_code:
        sb = p.paragraph_format.space_before
        if not sb or sb < Pt(8):
            p.paragraph_format.space_before = Pt(8)
    prev_code = is_code

doc.save(str(DOCX))
print(f"Post-processed: {DOCX}")
